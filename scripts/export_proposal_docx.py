"""Export the project proposal Markdown to the Word file opened in File Explorer."""

from __future__ import annotations

import re
import urllib.error
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "Project_Proposal_Web_Based_Citizen_Complaint_and_Service_Request.md"
DIAGRAM_DIR = ROOT / "scripts" / "assets" / "diagrams"
NAVY = RGBColor(0x00, 0x33, 0x66)
WORD_NAMES = [
    "ADAMA CITY CITIZEN COMPLIANT AND SERVICE REQUEST MANAGEMENT SYSTEM DOCUMENTATION.docx",
    "ADAMA CITY CITIZEN COMPLAINT MANAGEMENT SYSTEM DOCUMENTATION.docx",
    "NARRATIVE REPORT ON PRACTICAL ATTACHMENT - ADAMA STO.docx",
]
DEFAULT_ROMAN_HEADING = "Abstract (Executive Summary)"
DEFAULT_ARABIC_HEADING = "CHAPTER ONE"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003366")
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, fill="003366"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_formatted_runs(paragraph, text, size=12, color=None):
    pattern = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(10, size - 1), color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def add_heading(doc, text, level):
    sizes = {0: 22, 1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level else WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=16 if level else 6, after=8 if level else 4)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=NAVY)
    if level == 0:
        add_bottom_border(p)
    return p


def add_body(doc, text, italic=False, center=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, after=8)
    if italic:
        run = p.add_run(text)
        set_run_font(run, size=size, italic=True)
    else:
        add_formatted_runs(p, text, size=size)
    return p


def add_toc_entry(doc, title, page, level=0):
    """Academic TOC row: title, dotted leader, right-aligned page number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    before = 6 if level == 0 and title.upper().startswith("CHAPTER") else 0
    set_paragraph_spacing(p, before=before, after=2, line=1.0)
    p.paragraph_format.left_indent = Inches(0.4 * level)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    is_chapter = title.upper().startswith("CHAPTER")
    run = p.add_run(title)
    set_run_font(run, size=12, bold=is_chapter, color=RGBColor(0, 0, 0))
    tab = p.add_run("\t")
    set_run_font(tab, size=12, color=RGBColor(0, 0, 0))
    num = p.add_run(page)
    set_run_font(num, size=12, bold=is_chapter, color=RGBColor(0, 0, 0))
    return p


TOC_LINE = re.compile(r"^(\s*)(.+?)\s+\.{2,}\s+(\S+)\s*$")


def add_list_item(doc, text, ordered=False, number=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    set_paragraph_spacing(p, after=4)
    prefix = f"{number}. " if ordered else "• "
    run = p.add_run(prefix)
    set_run_font(run, size=12)
    add_formatted_runs(p, text, size=12)
    return p


def add_code_block(doc, text):
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        set_paragraph_spacing(p, after=0, line=1.0)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2, line=1.0)
            value = row[c_idx] if c_idx < len(row) else ""
            if r_idx == 0:
                shade_cell(cell, "003366")
                add_formatted_runs(p, value, size=10, color=RGBColor(255, 255, 255))
                for run in p.runs:
                    run.bold = True
            else:
                add_formatted_runs(p, value, size=10)
    doc.add_paragraph()


def add_image(doc, path: Path, caption: str | None = None):
    if not path.exists():
        add_body(doc, f"[Missing image: {path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    is_logo = path.name == "haramaya-university-logo.png"
    set_paragraph_spacing(p, before=2 if is_logo else 8, after=4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(1.7 if is_logo else 6.2))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, after=10)
        run = cap.add_run(caption)
        set_run_font(run, size=10, italic=True)


def render_mermaid(source: str, index: int) -> Path | None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    out = DIAGRAM_DIR / f"figure-{index:02d}.png"
    cleaned = re.sub(r"%%\{.*?\}%%", "", source, flags=re.S).strip()
    try:
        req = urllib.request.Request(
            "https://kroki.io/mermaid/png",
            data=cleaned.encode("utf-8"),
            headers={"Content-Type": "text/plain", "Accept": "image/png"},
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=45) as resp:
            data = resp.read()
        if data[:8] == b"\x89PNG\r\n\x1a\n":
            out.write_bytes(data)
            return out
    except (urllib.error.URLError, TimeoutError, OSError):
        return None
    return None


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-{3,}", line.strip()))


def add_page_field(run):
    """Insert a Word PAGE field so the footer number updates in Word."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)


def set_page_numbering(section, fmt="decimal", start=1):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))


def setup_centered_page_footer(section, fmt="decimal", start=1):
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
    run = paragraph.add_run()
    set_run_font(run, size=11, color=NAVY)
    add_page_field(run)
    set_page_numbering(section, fmt=fmt, start=start)
    section.footer_distance = Cm(1.25)


def start_new_numbered_section(doc, fmt, start=1):
    previous = doc.sections[-1]
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = previous.page_width
    section.page_height = previous.page_height
    section.left_margin = previous.left_margin
    section.right_margin = previous.right_margin
    section.top_margin = previous.top_margin
    section.bottom_margin = previous.bottom_margin
    section.different_first_page_header_footer = False
    setup_centered_page_footer(section, fmt=fmt, start=start)
    return section


def enable_update_fields_on_open(doc):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def convert(
    md_path: Path | None = None,
    word_names: list[str] | None = None,
    roman_heading: str = DEFAULT_ROMAN_HEADING,
    arabic_heading: str = DEFAULT_ARABIC_HEADING,
    extra_fallbacks: list[Path] | None = None,
):
    source = Path(md_path) if md_path else MD_PATH
    output_names = word_names if word_names else WORD_NAMES
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    cover_footer = section.first_page_footer
    cover_footer.is_linked_to_previous = False
    cover_p = cover_footer.paragraphs[0] if cover_footer.paragraphs else cover_footer.add_paragraph()
    cover_p.text = ""
    cover_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(cover_p, before=0, after=0, line=1.0)
    cover_run = cover_p.add_run("Submission date: 24 August 2026")
    set_run_font(cover_run, size=12, italic=True, color=NAVY)
    title_footer = section.footer
    title_footer.is_linked_to_previous = False
    if title_footer.paragraphs:
        title_footer.paragraphs[0].text = ""

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    i = 0
    mermaid_index = 0
    pending_caption = None
    front_matter_started = False
    body_started = False
    in_toc = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if in_toc:
            if stripped == "---" or stripped.startswith("#"):
                in_toc = False
            else:
                toc = TOC_LINE.match(line.rstrip())
                if toc:
                    leading = len(toc.group(1).replace("\t", "    "))
                    level = 0 if leading < 4 else 1 if leading < 8 else 2
                    add_toc_entry(doc, toc.group(2).strip(), toc.group(3).strip(), level)
                    i += 1
                    continue
                in_toc = False

        if stripped in ("<!-- pagebreak -->", "<!--pagebreak-->"):
            doc.add_page_break()
            i += 1
            continue

        if stripped.startswith("{center}"):
            add_body(doc, stripped[len("{center}"):].strip(), center=True)
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=4, after=8)
            add_bottom_border(p)
            i += 1
            continue

        if stripped.startswith("```"):
            fence = stripped[3:].strip()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            content = "\n".join(block)
            if fence.lower() == "mermaid":
                mermaid_index += 1
                image = render_mermaid(content, mermaid_index)
                caption = None
                if i < len(lines) and lines[i].strip().startswith("*Figure"):
                    caption = lines[i].strip().strip("*")
                    i += 1
                if image:
                    add_image(doc, image, caption)
                else:
                    add_body(doc, f"Figure diagram {mermaid_index} (see source Markdown if image did not render).", italic=True)
                    add_code_block(doc, content)
                    if caption:
                        add_body(doc, caption, italic=True, center=True, size=10)
            else:
                add_code_block(doc, content)
            continue

        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if img:
            rel = img.group(2).strip()
            path = (ROOT / rel).resolve()
            i += 1
            caption = None
            if i < len(lines) and lines[i].strip().startswith("*Figure"):
                caption = lines[i].strip().strip("*")
                i += 1
            add_image(doc, path, caption)
            continue

        if stripped.startswith("*Figure") and stripped.endswith("*"):
            add_body(doc, stripped.strip("*"), italic=True, center=True, size=10)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1].strip()):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            title = heading.group(2).strip()
            if title == roman_heading and not front_matter_started:
                start_new_numbered_section(doc, fmt="lowerRoman", start=1)
                front_matter_started = True
            elif title.startswith(arabic_heading) and not body_started:
                start_new_numbered_section(doc, fmt="decimal", start=1)
                body_started = True
            level = len(heading.group(1)) - 1
            add_heading(doc, title, level)
            if title.lower() == "table of contents":
                in_toc = True
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            add_list_item(doc, bullet.group(1))
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            add_list_item(doc, numbered.group(2), ordered=True, number=int(numbered.group(1)))
            i += 1
            continue

        add_body(doc, stripped)
        i += 1

    if pending_caption:
        add_body(doc, pending_caption, italic=True, center=True, size=10)

    enable_update_fields_on_open(doc)

    written = []
    errors = []
    fallbacks = extra_fallbacks or [
        ROOT / "ADAMA CITY CITIZEN COMPLAINT MANAGEMENT SYSTEM DOCUMENTATION - WITH PAGE NUMBERS.docx",
        ROOT / "DOCUMENTATION WITH PAGE NUMBERS.docx",
    ]
    for name in output_names:
        dest = ROOT / name
        candidates = [
            dest,
            ROOT / name.replace(".docx", " - UPDATED 23 AUG 2026.docx"),
            ROOT / name.replace(".docx", " - UPDATED 22 AUG 2026.docx"),
            *fallbacks,
        ]
        saved = False
        last_err = None
        for candidate in candidates:
            try:
                doc.save(str(candidate))
                written.append((candidate, candidate.stat().st_size, candidate.stat().st_mtime))
                if candidate != dest:
                    errors.append(f"{name} is open in Word. Saved: {candidate.name}")
                saved = True
                break
            except PermissionError as err:
                last_err = err
                continue
        if not saved:
            errors.append(f"Could not save {name}: {last_err}")

    print("WROTE")
    for path, size, mtime in written:
        print(f"{path.name} | {size} bytes")
    for err in errors:
        print("WARN", err)


if __name__ == "__main__":
    convert()
