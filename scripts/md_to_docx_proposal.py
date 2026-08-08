"""Convert the project proposal Markdown into a Word .docx with Mermaid diagrams as images."""
from __future__ import annotations

import base64
import re
import sys
import zlib
from pathlib import Path
from urllib.error import URLError, HTTPError
from urllib.request import Request, urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"c:\Users\HP\Desktop\Citizen")
MD_PATH = ROOT / "Project_Proposal_Web_Based_Citizen_Complaint_and_Service_Request.md"
DOCX_PATH = ROOT / "Project_Proposal_Web_Based_Citizen_Complaint_and_Service_Request.docx"
DOCX_FRONTEND = ROOT / "frontend" / "Project_Proposal_Web_Based_Citizen_Complaint_and_Service_Request.docx"
DOCX_FALLBACK = ROOT / "Project_Proposal_v2.11.docx"
DIAGRAM_DIR = ROOT / "scripts" / "_diagram_cache"
LOGO_PATH = ROOT / "scripts" / "assets" / "haramaya-logo.png"
RESULTS_DIR = ROOT / "scripts" / "assets" / "results"

DIAGRAM_CAPTIONS = {
    0: "Figure: System Architecture",
    1: "Figure: Use Case Diagram",
    2: "Figure: Entity-Relationship Diagram",
    3: "Figure: Data Flow Diagram (Level 0)",
    4: "Figure: System Workflow",
    5: "Figure: Project Timeline (Gantt Chart)",
    6: "Figure: Deployment Architecture",
}


def kroki_png(mermaid_source: str) -> bytes | None:
    compressed = zlib.compress(mermaid_source.encode("utf-8"), 9)
    encoded = base64.urlsafe_b64encode(compressed).decode("ascii")
    url = f"https://kroki.io/mermaid/png/{encoded}"
    req = Request(url, headers={"User-Agent": "CitizenProposalExport/1.0"})
    try:
        with urlopen(req, timeout=60) as resp:
            data = resp.read()
            if data[:8] == b"\x89PNG\r\n\x1a\n":
                return data
    except (URLError, HTTPError, TimeoutError) as err:
        print(f"  Kroki failed: {err}")
    return None


def mermaid_ink_png(mermaid_source: str) -> bytes | None:
    encoded = base64.urlsafe_b64encode(mermaid_source.encode("utf-8")).decode("ascii")
    url = f"https://mermaid.ink/img/{encoded}?type=png"
    req = Request(url, headers={"User-Agent": "CitizenProposalExport/1.0"})
    try:
        with urlopen(req, timeout=60) as resp:
            data = resp.read()
            if data[:8] == b"\x89PNG\r\n\x1a\n" or data[:2] == b"\xff\xd8":
                return data
    except (URLError, HTTPError, TimeoutError) as err:
        print(f"  mermaid.ink failed: {err}")
    return None


def normalize_diagram_png(data: bytes) -> bytes:
    """Flatten to white background and crop empty margins for clearer print layout."""
    try:
        from io import BytesIO
        from PIL import Image, ImageChops

        im = Image.open(BytesIO(data)).convert("RGBA")
        # Flatten transparency onto white (avoids dark/transparent canvases in Word)
        bg = Image.new("RGBA", im.size, (255, 255, 255, 255))
        flat = Image.alpha_composite(bg, im).convert("RGB")

        # Crop near-white / empty borders
        bg_ref = Image.new("RGB", flat.size, (255, 255, 255))
        diff = ImageChops.difference(flat, bg_ref)
        bbox = diff.getbbox()
        if bbox:
            pad = 12
            left = max(0, bbox[0] - pad)
            top = max(0, bbox[1] - pad)
            right = min(flat.width, bbox[2] + pad)
            bottom = min(flat.height, bbox[3] + pad)
            flat = flat.crop((left, top, right, bottom))

        # Upscale small diagrams so text stays readable in Word
        min_w = 1400
        if flat.width < min_w:
            scale = min_w / flat.width
            flat = flat.resize(
                (int(flat.width * scale), int(flat.height * scale)),
                Image.Resampling.LANCZOS,
            )

        buf = BytesIO()
        flat.save(buf, format="PNG", optimize=True)
        return buf.getvalue()
    except Exception as err:
        print(f"  PNG normalize skipped: {err}")
        return data


def _load_font(size: int, bold: bool = False):
    from PIL import ImageFont

    candidates = (
        [
            r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        ]
        if bold
        else [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\segoeui.ttf",
        ]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def render_use_case_png() -> bytes:
    """Draw a clear, print-friendly UML-style use case diagram (avoids Mermaid clutter)."""
    import math
    from io import BytesIO
    from PIL import Image, ImageDraw

    W, H = 1700, 900
    img = Image.new("RGB", (W, H), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    font_title = _load_font(20, bold=True)
    font_actor = _load_font(15, bold=True)
    font_uc = _load_font(14)
    font_inc = _load_font(11)
    navy = (0, 51, 102)
    ink = (25, 25, 25)
    line_c = (55, 55, 55)
    fill = (232, 238, 247)
    border = (0, 51, 102)

    # System boundary (tight around content — no large empty region)
    sys_box = (280, 55, 1420, 845)
    draw.rectangle(sys_box, outline=border, width=3)
    title = "Web-Based Citizen Complaint and Service Request Management System"
    tw = draw.textlength(title, font=font_title)
    tx = (sys_box[0] + sys_box[2] - tw) / 2
    draw.rectangle((tx - 8, 28, tx + tw + 8, 54), fill=(255, 255, 255))
    draw.text((tx, 32), title, fill=navy, font=font_title)

    def actor_points(cx: int, cy: int, label: str):
        head_r = 14
        draw.ellipse(
            (cx - head_r, cy - 48, cx + head_r, cy - 48 + 2 * head_r),
            outline=ink,
            width=2,
        )
        draw.line((cx, cy - 34, cx, cy), fill=ink, width=2)
        draw.line((cx - 20, cy - 20, cx + 20, cy - 20), fill=ink, width=2)
        draw.line((cx, cy, cx - 14, cy + 26), fill=ink, width=2)
        draw.line((cx, cy, cx + 14, cy + 26), fill=ink, width=2)
        lw = draw.textlength(label, font=font_actor)
        draw.text((cx - lw / 2, cy + 32), label, fill=ink, font=font_actor)
        return (cx, cy - 8)

    def oval_size(label: str):
        tw = draw.textlength(label, font=font_uc)
        return max(168, tw + 28), 40

    citizen_ucs = [
        "Register Account",
        "Login",
        "Submit Complaint",
        "Submit Service Request",
        "Track Submissions Status",
        "Update Profile",
        "View Notifications",
    ]
    admin_ucs = [
        "Manage Users",
        "Manage Departments",
        "Assign Submissions",
        "Update Status",
        "Generate Reports",
        "Monitor System Activities",
        "Search and Filter Submissions",
    ]
    officer_ucs = [
        "View Assigned Tasks",
        "Process Assigned Work",
        "Add Resolution Note",
    ]

    left_x, mid_x, right_x = 480, 850, 1220
    top_y, gap = 110, 95
    positions: dict[str, tuple[float, float]] = {}

    for i, name in enumerate(citizen_ucs):
        positions[name] = (left_x, top_y + i * gap)
    for i, name in enumerate(admin_ucs):
        positions[name] = (mid_x, top_y + i * gap)
    for i, name in enumerate(officer_ucs):
        positions[name] = (right_x, top_y + 120 + i * gap)

    cit = actor_points(120, 400, "Citizen")
    adm = actor_points(1580, 250, "Administrator")
    off = actor_points(1580, 620, "Department Officer")

    def edge_point(uc_name: str, from_left: bool):
        cx, cy = positions[uc_name]
        ow, _ = oval_size(uc_name)
        return (cx - ow / 2 if from_left else cx + ow / 2, cy)

    def draw_assoc(actor_pt, uc_name, from_left: bool):
        ux, uy = edge_point(uc_name, from_left=from_left)
        ax = actor_pt[0] + (22 if from_left else -22)
        draw.line((ax, actor_pt[1], ux, uy), fill=line_c, width=1)

    for name in citizen_ucs:
        draw_assoc(cit, name, from_left=True)
    for name in admin_ucs:
        draw_assoc(adm, name, from_left=False)
    for name in officer_ucs + ["Update Status", "View Notifications", "Login"]:
        draw_assoc(off, name, from_left=False)

    def draw_dashed(x1, y1, x2, y2, dash=8, gap=5):
        dx, dy = x2 - x1, y2 - y1
        dist = math.hypot(dx, dy) or 1
        ux, uy = dx / dist, dy / dist
        pos = 0.0
        draw_on = True
        while pos < dist:
            seg = dash if draw_on else gap
            end = min(pos + seg, dist)
            if draw_on:
                draw.line(
                    (x1 + ux * pos, y1 + uy * pos, x1 + ux * end, y1 + uy * end),
                    fill=(90, 90, 90),
                    width=1,
                )
            pos = end
            draw_on = not draw_on

    def include(src: str, dst: str, label_offset=(0, 10)):
        x1, y1 = positions[src]
        x2, y2 = positions[dst]
        draw_dashed(x1, y1 + 24, x2, y2 + 24)
        mx = (x1 + x2) / 2 + label_offset[0]
        my = (y1 + y2) / 2 + 24 + label_offset[1]
        label = "<<include>>"
        lw = draw.textlength(label, font=font_inc)
        draw.rectangle(
            (mx - lw / 2 - 2, my - 8, mx + lw / 2 + 2, my + 8),
            fill=(255, 255, 255),
        )
        draw.text((mx - lw / 2, my - 6), label, fill=(80, 80, 80), font=font_inc)

    include("Submit Complaint", "Login", label_offset=(80, 8))
    include("Submit Service Request", "Login", label_offset=(100, 8))
    include("Assign Submissions", "Update Status", label_offset=(0, 8))
    include("Process Assigned Work", "Update Status", label_offset=(50, 8))

    # Draw use-case ovals last so text is never covered by lines
    for name, (cx, cy) in positions.items():
        ow, oh = oval_size(name)
        box = (cx - ow / 2, cy - oh / 2, cx + ow / 2, cy + oh / 2)
        draw.ellipse(box, fill=fill, outline=border, width=2)
        tw = draw.textlength(name, font=font_uc)
        draw.text((cx - tw / 2, cy - 8), name, fill=ink, font=font_uc)

    buf = BytesIO()
    img.save(buf, format="PNG", optimize=True)
    return buf.getvalue()



def render_diagram(index: int, source: str) -> Path | None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    out = DIAGRAM_DIR / f"diagram_{index}.png"
    meta = DIAGRAM_DIR / f"diagram_{index}.src.hash"
    src_hash = str(zlib.crc32(source.encode("utf-8")) & 0xFFFFFFFF)
    if (
        out.exists()
        and out.stat().st_size > 500
        and meta.exists()
        and meta.read_text(encoding="utf-8").strip() == src_hash
    ):
        return out
    print(f"Rendering diagram {index}...")

    data = None
    # Use-case diagram: local high-clarity renderer (Mermaid layout is hard to read in Word)
    if index == 1:
        try:
            data = render_use_case_png()
            print("  Use-case diagram rendered locally")
        except Exception as err:
            print(f"  Local use-case render failed: {err}")

    if not data:
        data = kroki_png(source) or mermaid_ink_png(source)
        if data:
            data = normalize_diagram_png(data)

    if not data:
        # Keep previous image if online render fails (avoids blank figures in the document)
        if out.exists() and out.stat().st_size > 500:
            print(f"  Keeping previous {out.name}")
            return out
        return None

    out.write_bytes(data)
    meta.write_text(src_hash, encoding="utf-8")
    return out


def fit_image_width_inches(img_path: Path, max_width: float = 6.0, max_height: float = 7.2) -> float:
    """Scale image to fit page without overflowing (avoids blank pages from oversized figures)."""
    try:
        from PIL import Image

        w, h = Image.open(img_path).size
        if w <= 0 or h <= 0:
            return max_width
        width = max_width
        height = width * (h / w)
        if height > max_height:
            width = max_height * (w / h)
        return max(2.2, min(width, max_width))
    except Exception:
        return max_width


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    # Avoid orphan headings / forced page breaks that leave blank pages
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level <= 2 else 6)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    pf.page_break_before = False
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x00, 0x28, 0x68)
    return p


def add_paragraph_with_inline(doc: Document, text: str, style=None):
    p = doc.add_paragraph(style=style)
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos : m.start()])
            set_run_font(run)
        token = m.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("*"):
            run = p.add_run(token[1:-1])
            set_run_font(run)
            run.italic = True
        elif token.startswith("`"):
            run = p.add_run(token[1:-1])
            set_run_font(run, size=10)
            run.font.name = "Consolas"
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = p.add_run(label)
            set_run_font(run)
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_run_font(run)
    return p


def add_table(doc: Document, rows: list[list[str]]):
    if len(rows) < 1:
        return
    body_rows = rows
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", c.strip() or "-") for c in rows[1]):
        body_rows = [rows[0]] + rows[2:]

    cols = max(len(r) for r in body_rows)
    table = doc.add_table(rows=len(body_rows), cols=cols)
    table.style = "Table Grid"
    for ridx, row in enumerate(body_rows):
        for cidx in range(cols):
            cell = table.cell(ridx, cidx)
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[cidx].strip() if cidx < len(row) else ""
            val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
            val = re.sub(r"`([^`]+)`", r"\1", val)
            run = p.add_run(val)
            set_run_font(run, size=9, bold=(ridx == 0))
            if ridx == 0:
                run.font.color.rgb = RGBColor(0x00, 0x28, 0x68)
    doc.add_paragraph()


def add_page_number_field(paragraph):
    """Insert a Word PAGE field into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_page_number_format(section, fmt: str = "decimal", start: int | None = None):
    """fmt: 'lowerRoman' | 'upperRoman' | 'decimal'"""
    sect_pr = section._sectPr
    for el in sect_pr.xpath("./w:pgNumType"):
        sect_pr.remove(el)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    sect_pr.append(pg_num)


def setup_footer(section, linked_to_previous: bool = False):
    footer = section.footer
    footer.is_linked_to_previous = linked_to_previous
    if footer.paragraphs:
        p = footer.paragraphs[0]
        p.clear()
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(p)


def apply_section_margins(section):
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_centered(doc: Document, text: str, size=12, bold=False, space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or RGBColor(0x00, 0x28, 0x68))
    return p


TOC_ENTRY_RE = re.compile(
    r"^(?P<indent>\s*)(?P<title>.+?)\s*\.{2,}\s*(?P<page>[ivxlcdmIVXLCDM0-9]+)\s*$"
)


def toc_level_for(title: str, indent: str) -> int:
    """0 = front matter / chapter; 1 = x.y; 2 = x.y.z — also use leading spaces."""
    t = title.strip()
    if re.match(r"^\d+\.\d+\.\d+", t):
        return 2
    if re.match(r"^\d+\.\d+", t):
        return 1
    # Fallback from indentation (4 spaces ≈ one level)
    spaces = len(indent.replace("\t", "    "))
    if spaces >= 8:
        return 2
    if spaces >= 4:
        return 1
    return 0


def add_toc_entry(doc: Document, title: str, page: str, level: int = 0):
    """Academic TOC row: title .... page (dot leaders + right-aligned page)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    # Indent subsections like the sample (≈0.5" / 1.0")
    left_in = {0: 0.0, 1: 0.35, 2: 0.7}.get(level, 0.0)
    pf.left_indent = Inches(left_in)
    # Right-aligned tab with dotted leader (content width ≈ 6.2")
    tab_pos = Inches(6.2)
    pf.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    is_chapter = bool(re.match(r"^CHAPTER\s+", title, re.I))
    run = p.add_run(title.strip())
    set_run_font(run, size=11, bold=is_chapter)
    # Times-like academic look for TOC body text
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p.add_run("\t")
    page_run = p.add_run(page.strip())
    set_run_font(page_run, size=11, bold=is_chapter)
    page_run.font.name = "Times New Roman"
    page_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_cover_page(doc: Document):
    """Formal Haramaya University cover page (first page of preface)."""
    add_centered(doc, "HARAMAYA UNIVERSITY", size=16, bold=True, space_after=4)
    add_centered(
        doc,
        "COLLEGE OF COMPUTING AND INFORMATICS",
        size=13,
        bold=True,
        space_after=2,
    )
    add_centered(
        doc,
        "DEPARTMENT OF INFORMATION SCIENCE",
        size=12,
        bold=True,
        space_after=10,
    )

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.55))

    add_centered(doc, "PROJECT PROPOSAL ON", size=12, bold=True, space_before=6, space_after=8)
    add_centered(
        doc,
        "WEB-BASED CITIZEN COMPLAINT AND SERVICE REQUEST",
        size=13,
        bold=True,
        space_after=2,
    )
    add_centered(
        doc,
        "MANAGEMENT SYSTEM FOR ADAMA CITY ADMINISTRATION",
        size=13,
        bold=True,
        space_after=18,
    )

    add_centered(doc, "Prepared By:", size=12, bold=True, space_after=8)
    add_centered(doc, "Name                                                              ID", size=11, bold=True, space_after=4)
    add_centered(
        doc,
        "________________________________  ........................  ______________",
        size=11,
        space_after=16,
    )

    add_centered(
        doc,
        "Host Company: - Adama City Administration Science and Technology Office.",
        size=11,
        space_after=8,
    )
    add_centered(
        doc,
        "The Practical Attachment document submitted to Information Science Department.",
        size=11,
        space_after=16,
    )

    add_centered(doc, "Advisor: ________________________________", size=12, bold=True, space_after=6)
    add_centered(doc, "Submission Date: ____________________", size=11, space_after=6)
    add_centered(doc, "Haramaya, Ethiopia", size=12, bold=True, space_before=8, space_after=0)


def convert(md_text: str) -> Document:
    doc = Document()

    # Cover page section — no page numbers
    cover = doc.sections[0]
    apply_section_margins(cover)
    cover.footer.is_linked_to_previous = False
    if cover.footer.paragraphs:
        cover.footer.paragraphs[0].clear()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover_page(doc)

    # Preface (Declaration, Acknowledgement, TOC, Acronyms): Roman numerals i, ii, iii...
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_section_margins(front)
    set_page_number_format(front, fmt="lowerRoman", start=1)
    setup_footer(front, linked_to_previous=False)

    lines = md_text.splitlines()
    i = 0
    diagram_index = 0
    in_code = False
    code_lang = ""
    code_buf: list[str] = []
    body_started = False
    skip_title_block = True  # skip first H1/H2 title + metadata table (covered by cover page)

    while i < len(lines):
        line = lines[i]

        # Skip the opening title + metadata table (replaced by formal cover)
        if skip_title_block:
            if line.startswith("# ") or line.startswith("## Web-Based") or line.strip() == "---":
                i += 1
                continue
            if line.strip().startswith("|") and ("Field" in line or "Detail" in line or "**Project Title**" in line or "**Organization**" in line or "**Document Version**" in line or "**Date**" in line or "**Prepared By**" in line or "**Institution**" in line or "**Supervisor" in line or "**Host Organization**" in line or line.strip().startswith("|---")):
                i += 1
                continue
            if not line.strip():
                i += 1
                continue
            skip_title_block = False

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_buf = []
            else:
                code = "\n".join(code_buf)
                if code_lang == "mermaid":
                    caption = DIAGRAM_CAPTIONS.get(diagram_index, f"Figure {diagram_index + 1}")
                    img_path = render_diagram(diagram_index, code)
                    if img_path and img_path.exists():
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(2)
                        # Keep figure with caption only — avoid pushing large images to a blank page
                        p.paragraph_format.keep_with_next = True
                        run = p.add_run()
                        # Use-case diagram (index 1): wider + shorter height budget so it stays on one page
                        if diagram_index == 1:
                            width = fit_image_width_inches(
                                img_path, max_width=6.3, max_height=5.4
                            )
                        elif diagram_index == 2:
                            width = fit_image_width_inches(
                                img_path, max_width=5.6, max_height=6.6
                            )
                        else:
                            width = fit_image_width_inches(
                                img_path, max_width=5.8, max_height=6.5
                            )
                        run.add_picture(str(img_path), width=Inches(width))
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.paragraph_format.space_after = Pt(8)
                        cap.paragraph_format.keep_with_next = False
                        r = cap.add_run(caption)
                        set_run_font(r, size=10, bold=True, color=RGBColor(0x00, 0x28, 0x68))
                    else:
                        p = doc.add_paragraph()
                        r = p.add_run("[diagram could not be rendered online]")
                        set_run_font(r, size=10, color=RGBColor(0x99, 0x1B, 0x1B))
                        pre = doc.add_paragraph()
                        r2 = pre.add_run(code)
                        set_run_font(r2, size=8)
                        r2.font.name = "Consolas"
                    diagram_index += 1
                else:
                    pre = doc.add_paragraph()
                    r = pre.add_run(code)
                    set_run_font(r, size=8)
                    r.font.name = "Consolas"
                in_code = False
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            # Skip markdown rules — they create sparse / blank-looking pages in Word
            i += 1
            continue

        # Markdown image: ![alt](path)
        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if img_m:
            alt, rel = img_m.group(1), img_m.group(2).strip()
            img_path = (ROOT / rel).resolve() if not Path(rel).is_absolute() else Path(rel)
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True
                run = p.add_run()
                width = fit_image_width_inches(img_path, max_width=5.8, max_height=6.8)
                run.add_picture(str(img_path), width=Inches(width))
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[missing image: {rel}]")
                set_run_font(r, size=10, color=RGBColor(0x99, 0x1B, 0x1B))
            i += 1
            continue

        # Italic figure caption: *Figure ...*
        cap_m = re.match(r"^\*(Figure[^*]*)\*\s*$", line.strip())
        if cap_m:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(cap_m.group(1))
            set_run_font(r, size=10, bold=True, color=RGBColor(0x00, 0x28, 0x68))
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                cells = [c for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            text = re.sub(r"\s*\{#[^}]+\}\s*$", "", m.group(2).strip())
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)

            if not body_started and (
                re.match(r"^CHAPTER ONE:", text, re.I)
                or re.match(r"^1\.\s+Background", text)
            ):
                body = doc.add_section(WD_SECTION.NEW_PAGE)
                apply_section_margins(body)
                set_page_number_format(body, fmt="decimal", start=1)
                setup_footer(body, linked_to_previous=False)
                body_started = True

            add_heading_styled(doc, text, level)
            # Prefer academic title casing for TOC heading
            if re.match(r"^table of contents$", text, re.I):
                for run in doc.paragraphs[-1].runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            i += 1
            continue

        # Academic TOC rows: Title ........ page
        toc_m = TOC_ENTRY_RE.match(line.rstrip())
        if toc_m:
            title = toc_m.group("title").strip()
            page = toc_m.group("page").strip()
            level = toc_level_for(title, toc_m.group("indent"))
            add_toc_entry(doc, title, page, level)
            i += 1
            continue

        m = re.match(r"^[-*]\s+(.*)$", line)
        if m:
            add_paragraph_with_inline(doc, m.group(1), style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            add_paragraph_with_inline(doc, m.group(2), style="List Number")
            i += 1
            continue

        add_paragraph_with_inline(doc, line.strip())
        i += 1

    return doc


def main() -> int:
    if not MD_PATH.exists():
        print(f"Missing: {MD_PATH}", file=sys.stderr)
        return 1

    print("Converting Markdown to Word (.docx) with page numbers...")
    md = MD_PATH.read_text(encoding="utf-8")
    doc = convert(md)

    targets = [DOCX_PATH, DOCX_FRONTEND]
    saved = []
    for path in targets:
        try:
            doc.save(path)
            saved.append(path)
            print(f"Wrote {path} ({path.stat().st_size:,} bytes)")
        except PermissionError:
            print(f"Could not overwrite {path} (file may be open).")

    if not saved:
        doc.save(DOCX_FALLBACK)
        print(f"Wrote fallback {DOCX_FALLBACK} ({DOCX_FALLBACK.stat().st_size:,} bytes)")
        print("Close the open Word file and re-run to update the main .docx.")
    elif DOCX_PATH not in saved:
        try:
            doc.save(DOCX_FALLBACK)
            print(f"Also wrote {DOCX_FALLBACK}")
        except PermissionError:
            pass

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
